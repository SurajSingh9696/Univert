#!/usr/bin/env python3
"""
HTML to DOCX converter using BeautifulSoup and python-docx.
Usage: python html_to_docx.py <input_html> <output_docx>
"""

import sys
from pathlib import Path

def html_to_docx(html_path, docx_path):
    """Convert HTML to DOCX by extracting text and basic formatting."""
    try:
        from bs4 import BeautifulSoup
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Read HTML content
        with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        
        # Process HTML elements
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'div', 'span', 'pre', 'code']):
            text = element.get_text(strip=True)
            if not text:
                continue
                
            tag_name = element.name
            
            if tag_name in ['h1']:
                heading = doc.add_heading(text, level=1)
            elif tag_name in ['h2']:
                heading = doc.add_heading(text, level=2)
            elif tag_name in ['h3', 'h4', 'h5', 'h6']:
                heading = doc.add_heading(text, level=3)
            elif tag_name in ['pre', 'code']:
                # Code block - use monospace-like formatting
                para = doc.add_paragraph(text)
                para.style = 'No Spacing'
            elif tag_name == 'li':
                doc.add_paragraph(text, style='List Bullet')
            else:
                doc.add_paragraph(text)
        
        # If no content was extracted, try getting all text
        if len(doc.paragraphs) == 0:
            all_text = soup.get_text(separator='\n', strip=True)
            for line in all_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
        
        doc.save(docx_path)
        print(f"Successfully converted {html_path} to {docx_path}")
        return True
        
    except ImportError as e:
        print(f"Missing required library: {e}", file=sys.stderr)
        return False
    except Exception as e:
        print(f"Error converting HTML to DOCX: {e}", file=sys.stderr)
        return False

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python html_to_docx.py <input_html> <output_docx>", file=sys.stderr)
        sys.exit(1)
    
    input_html = sys.argv[1]
    output_docx = sys.argv[2]
    
    success = html_to_docx(input_html, output_docx)
    sys.exit(0 if success else 1)
